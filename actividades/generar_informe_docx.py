# -*- coding: utf-8 -*-
"""Genera el informe de resultados en formato .docx para anexar a la tesis."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CORAL = RGBColor(0xFF, 0x5C, 0x39)
TEAL  = RGBColor(0x0F, 0xB5, 0xBA)
BERRY = RGBColor(0xC8, 0x1D, 0x7A)
PLUM  = RGBColor(0x6B, 0x2D, 0x8E)
INK   = RGBColor(0x1B, 0x0E, 0x05)
MUTED = RGBColor(0x95, 0x7B, 0x5E)

doc = Document()

# Márgenes
for section in doc.sections:
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

# Estilos base
styles = doc.styles
normal = styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)

def add_heading(text, level=1, color=INK):
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    run.font.color.rgb = color
    run.font.name = 'Calibri'
    return h

def add_para(text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p

def add_kv_line(label, value):
    p = doc.add_paragraph()
    r1 = p.add_run(label + ': ')
    r1.bold = True
    p.add_run(value)
    return p

def add_callout(title, body, color=CORAL):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    r1 = p.add_run(title + ' ')
    r1.bold = True
    r1.font.color.rgb = color
    p.add_run(body)

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        para = hdr[i].paragraphs[0]
        run = para.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
    if col_widths:
        for r in table.rows:
            for c, w in zip(r.cells, col_widths):
                c.width = w
    return table

# =========================== PORTADA ===========================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('INFORME DE RESULTADOS')
run.bold = True
run.font.size = Pt(11)
run.font.color.rgb = MUTED

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Evaluación de la plataforma 3D de unidades de refinación\ny las presentaciones interactivas')
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = INK

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Trabajo de investigación · Maestría · UNIPAZ')
run.italic = True
run.font.size = Pt(12)
run.font.color.rgb = MUTED

doc.add_paragraph()
doc.add_paragraph()

# Ficha técnica
add_heading('Ficha técnica del estudio', level=2, color=CORAL)
add_kv_line('Instrumento', 'Encuesta estructurada tipo Likert (escala 1-5)')
add_kv_line('Población', 'Estudiantes de Ingeniería de Producción, 6° semestre')
add_kv_line('Tamaño muestral', 'n = 12 estudiantes')
add_kv_line('Fecha de aplicación', '05 de junio de 2026')
add_kv_line('Ítems totales', '25 (18 Likert + 5 perfil + 2 abiertas)')
add_kv_line('Dimensiones evaluadas', '3 (Plataforma 3D · Presentaciones · Aprendizaje percibido)')
add_kv_line('Curso', 'Procesos Industriales I')

doc.add_page_break()

# =========================== 1. RESUMEN EJECUTIVO ===========================
add_heading('1. Resumen ejecutivo', level=1, color=CORAL)
add_para(
    'La evaluación arroja valoraciones consistentemente altas en las tres dimensiones del '
    'instrumento. Los promedios superan ampliamente el punto medio teórico de la escala '
    '(M = 3.0), ubicándose en el rango interpretativo de "de acuerdo" a "totalmente de acuerdo".'
)

add_heading('Indicadores clave', level=2)
add_table(
    ['Dimensión', 'Media (M)', 'Desv. Estándar (DE)', 'α de Cronbach'],
    [
        ['Plataforma 3D',           '4.51', '1.03', '0.988'],
        ['Presentaciones',          '4.51', '0.94', '0.986'],
        ['Aprendizaje percibido',   '4.38', '1.08', '0.977'],
    ]
)

add_heading('Conclusiones clave', level=2)
for item in [
    'Las dos modalidades (3D y presentaciones) son percibidas como equivalentes en valor '
    '(M = 4.51 en ambas), lo que respalda la hipótesis de complementariedad del diseño didáctico.',
    'El ítem mejor valorado es la intención de recomendar el visor 3D (M = 4.75; el 83.3% '
    'otorgó la puntuación máxima).',
    'Los ítems con menor puntuación pertenecen a la dimensión Aprendizaje: comprensión '
    'de la relación entre unidades (M = 4.25) y comprensión general (M = 4.33), aún '
    'dentro del rango "muy de acuerdo".',
    'La consistencia interna de las escalas es excelente en las tres dimensiones '
    '(α > 0.97), con las consideraciones reportadas en el apartado de limitaciones.'
]:
    p = doc.add_paragraph(item, style='List Number')

doc.add_page_break()

# =========================== 2. METODOLOGÍA ===========================
add_heading('2. Metodología', level=1, color=TEAL)
add_para(
    'El estudio empleó un diseño descriptivo de corte transversal con enfoque cuantitativo '
    'y un componente cualitativo complementario para las preguntas abiertas.'
)

add_heading('2.1 Instrumento', level=2)
for item in [
    '25 ítems organizados en cinco bloques: Perfil (5), Plataforma 3D (7), '
    'Presentaciones (7), Aprendizaje percibido (4) y Comentarios abiertos (2).',
    'Escala Likert de 5 niveles en los bloques 2-4: 1 = Totalmente en desacuerdo; '
    '5 = Totalmente de acuerdo.',
    'Aplicación digital autoadministrada vía formulario HTML responsive; '
    'almacenamiento en base de datos PostgreSQL (Supabase).'
]:
    doc.add_paragraph(item, style='List Bullet')

add_heading('2.2 Procedimiento ético', level=2)
for item in [
    'Consentimiento informado al inicio del instrumento.',
    'Tratamiento de datos personales conforme a la Ley 1581 de 2012 (Habeas Data, Colombia).',
    'Participación voluntaria y respuestas tratadas exclusivamente de forma agregada.'
]:
    doc.add_paragraph(item, style='List Bullet')

add_heading('2.3 Procesamiento de datos', level=2)
for item in [
    'Cálculo de estadísticos descriptivos (M, Mdn, Mo, DE, mín, máx) por ítem y por dimensión.',
    'Análisis de confiabilidad mediante el coeficiente Alfa de Cronbach para cada dimensión.',
    'Codificación temática inductiva para las respuestas a las preguntas abiertas.'
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# =========================== 3. CARACTERIZACIÓN ===========================
add_heading('3. Caracterización de la muestra', level=1, color=BERRY)
add_para(
    'La muestra está conformada por 12 estudiantes del programa de Ingeniería de Producción, '
    'todos cursando sexto semestre del curso Procesos Industriales I.'
)

add_table(
    ['Variable', 'Categoría', 'n', '%'],
    [
        ['Género',                              'Femenino',           '5', '41.7%'],
        ['',                                    'Masculino',          '7', '58.3%'],
        ['Experiencia con 3D',                  'No, primera vez',    '7', '58.3%'],
        ['',                                    'Sí, alguna vez',     '3', '25.0%'],
        ['',                                    'Sí, varias veces',   '2', '16.7%'],
        ['Conocimiento previo en refinación',   'Básico',             '7', '58.3%'],
        ['',                                    'Intermedio',         '5', '41.7%'],
    ]
)

doc.add_page_break()

# =========================== 4. ANÁLISIS DESCRIPTIVO POR DIMENSIÓN ===========================
add_heading('4. Análisis descriptivo por dimensión', level=1, color=CORAL)
add_para(
    'Se presentan medias, mediana, desviación estándar y coeficiente Alfa de Cronbach '
    'para cada una de las tres dimensiones evaluadas.'
)

add_table(
    ['Dimensión', 'k (ítems)', 'M', 'Mdn', 'DE', 'α de Cronbach'],
    [
        ['Plataforma 3D',                '7', '4.51', '5.00', '1.03', '0.988'],
        ['Presentaciones interactivas',  '7', '4.51', '4.93', '0.94', '0.986'],
        ['Aprendizaje percibido',        '4', '4.38', '4.88', '1.08', '0.977'],
    ]
)

add_callout(
    'Interpretación de α:',
    'los tres valores superan el umbral de confiabilidad excelente (α ≥ 0.90). '
    'Esto sugiere que los ítems de cada dimensión miden un constructo coherente. '
    'Ver consideraciones en el apartado de limitaciones.',
    TEAL
)

doc.add_page_break()

# =========================== 5. ANÁLISIS POR ÍTEM ===========================
add_heading('5. Análisis estadístico ítem por ítem', level=1, color=BERRY)

add_heading('5.1 Dimensión 1: Plataforma 3D', level=2)
add_table(
    ['Ítem', 'M', 'Mdn', 'Mo', 'DE', 'Mín', 'Máx'],
    [
        ['06. Fácil aprender a usar',           '4.50', '5', '5', '1.17', '1', '5'],
        ['07. Interfaz clara',                  '4.50', '5', '5', '1.17', '1', '5'],
        ['08. Funciona con fluidez',            '4.42', '5', '5', '1.16', '1', '5'],
        ['09. Modelos representan con realismo','4.42', '5', '5', '1.16', '1', '5'],
        ['10. Etiquetas son útiles',            '4.50', '5', '5', '1.17', '1', '5'],
        ['11. Motivó a aprender más',           '4.50', '5', '5', '0.90', '2', '5'],
        ['12. Recomendaría a otros',            '4.75', '5', '5', '0.62', '3', '5'],
    ]
)

add_heading('5.2 Dimensión 2: Presentaciones interactivas', level=2)
add_table(
    ['Ítem', 'M', 'Mdn', 'Mo', 'DE', 'Mín', 'Máx'],
    [
        ['13. Explicaciones claras',            '4.42', '5', '5', '0.90', '2', '5'],
        ['14. Animaciones ayudan a entender',   '4.67', '5', '5', '0.89', '2', '5'],
        ['15. Diseño visual atractivo',         '4.58', '5', '5', '1.16', '1', '5'],
        ['16. Ritmo adecuado',                  '4.50', '5', '5', '0.90', '2', '5'],
        ['17. Funciona en mobile',              '4.42', '5', '5', '1.16', '1', '5'],
        ['18. Complementan al 3D',              '4.50', '5', '5', '0.90', '2', '5'],
        ['19. Volvería a consultar',            '4.50', '5', '5', '0.90', '2', '5'],
    ]
)

add_heading('5.3 Dimensión 3: Aprendizaje percibido', level=2)
add_table(
    ['Ítem', 'M', 'Mdn', 'Mo', 'DE', 'Mín', 'Máx'],
    [
        ['20. Comprendo mejor el flujo general',     '4.33', '4.5', '5', '0.89', '2', '5'],
        ['21. Identifico la relación entre unidades','4.25', '5',   '5', '1.22', '1', '5'],
        ['22. Más efectiva que método tradicional',  '4.42', '5',   '5', '1.16', '1', '5'],
        ['23. Recomendaría incorporarlo al curso',   '4.50', '5',   '5', '1.17', '1', '5'],
    ]
)

doc.add_page_break()

# =========================== 6. DISTRIBUCIÓN ===========================
add_heading('6. Distribución de frecuencias por nivel Likert', level=1, color=PLUM)
add_para(
    'El patrón de respuestas muestra una fuerte asimetría positiva: más del 75% de las '
    'respuestas se concentran en los niveles 4 ("de acuerdo") y 5 ("totalmente de acuerdo") '
    'en las tres dimensiones.'
)

add_table(
    ['Dimensión', '1 (TD)', '2 (D)', '3 (N)', '4 (A)', '5 (TA)', 'Total resp.'],
    [
        ['Plataforma 3D',   '4 (4.8%)',  '1 (1.2%)', '1 (1.2%)', '16 (19.0%)', '62 (73.8%)', '84'],
        ['Presentaciones',  '2 (2.4%)',  '5 (6.0%)', '0 (0.0%)', '18 (21.4%)', '59 (70.2%)', '84'],
        ['Aprendizaje',     '3 (6.3%)',  '1 (2.1%)', '1 (2.1%)', '10 (20.8%)', '33 (68.8%)', '48'],
    ]
)
add_para('Nota: TD = Totalmente en desacuerdo · D = En desacuerdo · N = Neutral · A = De acuerdo · TA = Totalmente de acuerdo.', italic=True, size=9)

doc.add_page_break()

# =========================== 7. ANÁLISIS CUALITATIVO ===========================
add_heading('7. Análisis cualitativo de respuestas abiertas', level=1, color=TEAL)
add_para(
    'Se aplicó codificación temática inductiva a las respuestas de los ítems 24 ("lo que '
    'más gustó") y 25 ("mejoras"). Las categorías emergieron del análisis del contenido '
    'de las respuestas.'
)

add_heading('7.1 Pregunta 24 · ¿Qué fue lo que más te gustó?', level=2)
add_para('Categorías emergentes y frecuencia de menciones:', bold=True)
add_table(
    ['Categoría temática', 'Menciones'],
    [
        ['Visualización gráfica del proceso',                '8'],
        ['Comprensión y aprendizaje activo',                 '6'],
        ['Interactividad y manipulación de modelos',         '3'],
        ['Diseño visual y facilidad de uso',                 '3'],
        ['Acceso a equipos no disponibles físicamente',      '2'],
    ]
)
add_para('Citas representativas:', bold=True)
for src, text in [
    ('E08 (Femenino, Intermedio)',
     'Es una alternativa excelente para nosotros como estudiantes ya que es difícil '
     'encontrar ese tipo de interacciones más que todo para una clase; ayudan a entender '
     'más ya que muchos no hemos visto un equipo en la vida real.'),
    ('E09 (Masculino, Básico)',
     'Como es una herramienta que puedes manipular como Solidworks, permite que te '
     'concentres en puntos concretos a tu ritmo.'),
    ('E07 (Masculino, Básico)',
     'Ver un prototipo de una unidad atmosférica y torres de destilación y entre otras, '
     'donde se veía detalladamente el funcionamiento de cada una.'),
    ('E01 (Femenino, Básico)',
     'El aprendizaje con ejemplo; muchas veces lo gráfico y con herramientas que podemos '
     'manejar alienta mucho el conocimiento.'),
]:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    r = p.add_run(src + ': ')
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = BERRY
    r2 = p.add_run('«' + text + '»')
    r2.italic = True
    r2.font.size = Pt(10)

add_heading('7.2 Pregunta 25 · ¿Qué mejorarías?', level=2)
add_para('Categorías emergentes y frecuencia de menciones:', bold=True)
add_table(
    ['Categoría temática', 'Menciones'],
    [
        ['Sin sugerencias (satisfacción total)',           '5'],
        ['Rendimiento / velocidad de carga',               '2'],
        ['Mayor interactividad / construcción guiada',     '2'],
        ['Trazabilidad del flujo del proceso',             '2'],
        ['Compatibilidad de navegador',                    '1'],
    ]
)
add_para('Citas representativas:', bold=True)
for src, text in [
    ('E10 (Femenino, Básico)',
     'Que tuviera un paso a paso de dónde se inicia el procedimiento, para así entender '
     'más de dónde sale y hasta dónde termina.'),
    ('E03 (Masculino, Intermedio)',
     'Que pudiéramos ver desde cero y hacerlo nosotros.'),
    ('E09 (Masculino, Básico)',
     'Mostrar cómo circulan los componentes, como el crudo; de resto todo bien.'),
    ('E06 (Masculino, Intermedio)',
     'Que las presentaciones tengan menos texto y más desarrollo del proceso.'),
]:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    r = p.add_run(src + ': ')
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = BERRY
    r2 = p.add_run('«' + text + '»')
    r2.italic = True
    r2.font.size = Pt(10)

add_callout(
    'Insight metodológico:',
    'existe convergencia entre los datos cuantitativos y los cualitativos. El bloque '
    '"Aprendizaje" obtuvo la puntuación más baja (M = 4.38) y específicamente el ítem 21 '
    '("identifico la relación entre unidades", M = 4.25). Los comentarios refuerzan esto '
    'al solicitar trazabilidad del flujo entre unidades. Esta es una recomendación de '
    'mejora respaldada por triangulación de fuentes.',
    TEAL
)

doc.add_page_break()

# =========================== 8. TRIANGULACIÓN ===========================
add_heading('8. Triangulación: diferencias por experiencia previa', level=1, color=CORAL)
add_para(
    'Comparación de las medias por dimensión según el grado de experiencia previa con '
    'herramientas 3D. Dado el tamaño muestral (n = 12), se reporta a nivel descriptivo, '
    'no inferencial.'
)
add_table(
    ['Experiencia previa', 'n', 'M · 3D', 'M · Presentaciones', 'M · Aprendizaje'],
    [
        ['No, primera vez',   '7', '4.96', '4.84', '4.71'],
        ['Sí, alguna vez',    '3', '4.62', '4.62', '4.58'],
        ['Sí, varias veces',  '2', '2.79', '3.07', '2.63'],
    ]
)
add_callout(
    'Observación importante:',
    'el subgrupo con "varias experiencias previas" (n = 2) presenta promedios marcadamente '
    'más bajos. Al inspeccionar los datos individuales, se identifica que un caso (E06) '
    'reportó valoraciones extremas (1-2) en casi todos los ítems mientras simultáneamente '
    'emitió comentarios positivos en las preguntas abiertas. Este patrón sugiere posible '
    'respuesta atípica (extreme low responder) y debe documentarse en limitaciones.',
    BERRY
)

doc.add_page_break()

# =========================== 9. LIMITACIONES ===========================
add_heading('9. Limitaciones del estudio', level=1, color=RGBColor(0xD6, 0x28, 0x28))

for i, (title, body) in enumerate([
    ('Tamaño muestral reducido (n = 12).',
     'Los hallazgos deben interpretarse como evaluación piloto. No es posible aplicar '
     'pruebas inferenciales con potencia estadística suficiente. Los coeficientes Alfa de '
     'Cronbach calculados son referenciales; se recomienda replicar el estudio con n ≥ 30 '
     'para obtener valores definitivos.'),
    ('Alfa inflado por respuesta atípica.',
     'La presencia de un respondiente (E06) con puntuaciones extremas bajas distorsiona al '
     'alza el cálculo de α al aumentar la varianza total de la escala. Recalcular '
     'excluyendo este caso podría dar valores α más realistas; se recomienda análisis de '
     'sensibilidad para futuras réplicas.'),
    ('Sesgo de aceptabilidad social.',
     'La aplicación en contexto académico, con la docente como investigadora, puede haber '
     'introducido un sesgo hacia respuestas positivas. Trabajos futuros pueden incorporar '
     'evaluadores externos o aplicación en ausencia del docente.'),
    ('Diseño descriptivo de corte transversal.',
     'No permite establecer relaciones causales entre el uso de la plataforma y el '
     'aprendizaje real. Se sugiere un diseño cuasi-experimental con grupo control y '
     'mediciones pre-test/post-test en estudios posteriores.'),
    ('Una sola aplicación.',
     'No se evalúa el posible efecto de novedad. Se recomienda re-aplicar la encuesta '
     'tras varios usos para discriminar la evaluación de la herramienta del entusiasmo '
     'inicial.'),
], start=1):
    p = doc.add_paragraph(style='List Number')
    r = p.add_run(title + ' ')
    r.bold = True
    p.add_run(body)

doc.add_page_break()

# =========================== 10. CONCLUSIONES ===========================
add_heading('10. Conclusiones de la evaluación piloto', level=1, color=RGBColor(0x2B, 0xAE, 0x66))

for item in [
    'La plataforma 3D y las presentaciones interactivas obtienen valoraciones altamente '
    'positivas en los estudiantes de Procesos Industriales I, con medias por dimensión '
    'entre 4.38 y 4.51 sobre una escala de 5.',
    'La hipótesis de complementariedad de ambos recursos queda respaldada por la '
    'equivalencia de medias entre las dimensiones 3D y Presentaciones (M = 4.51 en ambas) '
    'y por la valoración explícita del ítem 18 ("complementan al 3D", M = 4.50).',
    'El instrumento presenta excelente consistencia interna en condiciones piloto '
    '(α > 0.97 en las tres dimensiones), lo que respalda su uso en estudios de mayor '
    'escala con la salvedad indicada en limitaciones.',
    'La triangulación cuanti-cualitativa identifica la trazabilidad del flujo entre '
    'unidades como la principal oportunidad de mejora del diseño didáctico, dato que '
    'aparece tanto en el ítem 21 (M = 4.25) como en los comentarios abiertos.',
    'Para la siguiente fase del estudio se recomienda: (a) ampliar la muestra a n ≥ 30; '
    '(b) implementar las mejoras sugeridas por los estudiantes; (c) aplicar un diseño '
    'cuasi-experimental con grupo control; (d) incluir medición objetiva de aprendizaje '
    '(no solo percibido) mediante un instrumento de conocimiento.',
]:
    doc.add_paragraph(item, style='List Number')

doc.add_page_break()

# =========================== ANEXO ===========================
add_heading('Anexo · Matriz de respuestas individuales', level=1, color=MUTED)
add_para('Datos crudos por estudiante (identificadores anonimizados E01-E12), promedios por dimensión.')

add_table(
    ['ID', 'Género', 'Experiencia previa', 'Nivel previo', 'M · 3D', 'M · Pres', 'M · Apr'],
    [
        ['E01', 'F', 'Primera vez',  'Básico',     '5.00', '5.00', '5.00'],
        ['E02', 'M', 'Alguna vez',   'Intermedio', '5.00', '5.00', '5.00'],
        ['E03', 'M', 'Varias veces', 'Intermedio', '4.14', '4.43', '4.00'],
        ['E04', 'F', 'Primera vez',  'Intermedio', '5.00', '4.86', '4.50'],
        ['E05', 'M', 'Alguna vez',   'Básico',     '4.00', '4.00', '4.25'],
        ['E06', 'M', 'Varias veces', 'Intermedio', '1.43', '1.71', '1.25'],
        ['E07', 'M', 'Primera vez',  'Básico',     '5.00', '4.86', '5.00'],
        ['E08', 'F', 'Primera vez',  'Intermedio', '5.00', '5.00', '5.00'],
        ['E09', 'M', 'Alguna vez',   'Básico',     '4.86', '5.00', '4.75'],
        ['E10', 'F', 'Primera vez',  'Básico',     '4.71', '4.29', '3.75'],
        ['E11', 'M', 'Primera vez',  'Básico',     '5.00', '5.00', '5.00'],
        ['E12', 'F', 'Primera vez',  'Básico',     '5.00', '5.00', '5.00'],
    ]
)
add_para('Nota: la fila E06 corresponde al respondiente con patrón atípico identificado en la sección de limitaciones (sección 9.2).', italic=True, size=9)

# Guardar
import os
out_path = os.path.join(os.path.dirname(__file__), 'informe_encuesta.docx')
doc.save(out_path)
print(f'OK: {out_path}')
